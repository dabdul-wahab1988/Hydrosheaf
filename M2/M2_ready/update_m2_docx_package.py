from __future__ import annotations

import io
import zipfile
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
READY = ROOT / "M2" / "M2_ready"
FIG_DIR = ROOT / "M2" / "m2_benchmark" / "figures" / "Manuscript_Ready"


def replace_paragraphs(doc: Document, mapping: dict[int, str]) -> None:
    for idx, text in mapping.items():
        p = doc.paragraphs[idx]
        for run in p.runs:
            run.text = ""
        if p.runs:
            p.runs[0].text = text
        else:
            p.add_run(text)


def set_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            if c < len(table.columns):
                table.cell(r, c).text = value
    for r in range(len(rows), len(table.rows)):
        for c in range(len(table.columns)):
            table.cell(r, c).text = ""


def replace_image(doc: Document, shape_index: int, image_bytes: bytes, alt: str) -> None:
    shape = doc.inline_shapes[shape_index]
    rid = shape._inline.graphic.graphicData.pic.blipFill.blip.get(qn("r:embed"))
    doc.part.rels[rid].target_part._blob = image_bytes
    shape._inline.docPr.set("descr", alt)
    shape._inline.docPr.set("title", alt)


def image_bytes(path: Path) -> bytes:
    return path.read_bytes()


def font(size: int, bold: bool = False):
    name = "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"
    return ImageFont.truetype(name, size=size)


def make_updated_graphical_abstract() -> bytes:
    source = READY / "Graphical abstract-CG.docx"
    with zipfile.ZipFile(source) as z:
        media = [n for n in z.namelist() if n.startswith("word/media/") and n.endswith(".png")][0]
        source_bytes = z.read(media)
    im = Image.open(io.BytesIO(source_bytes)).convert("RGB")
    d = ImageDraw.Draw(im)
    # Replace stale headline/claim text while retaining the established layout.
    d.rectangle((150, 46, 1030, 101), fill="#f7fbfd")
    d.text((180, 57), "Sparse observations → age-constrained topology → evidence-tiered process networks", font=font(18), fill="#526273")
    d.rectangle((1030, 42, 1390, 102), fill="#f7fbfd")
    d.text((1050, 48), "Evidence-tiered process", font=font(23, True), fill="#123f63")
    d.text((1050, 76), "screening outputs", font=font(23, True), fill="#123f63")

    d.rounded_rectangle((1060, 190, 1370, 348), radius=12, fill="#ffffff", outline="#c8d7e3", width=2)
    d.text((1080, 204), "Benchmark evidence", font=font(18, True), fill="#123f63")
    lines = [
        "• Chemistry R² = 0.955; reaction R² = 0.57",
        "• No-prior topology F1 = 0.618; prior F1 = 1.00",
        "• Synthetic age R² = 0.98; public parity R² = 0.71",
        "• PHREEQC-compatible proxy feasible = 89.1%",
        "• Ghana: 208 edges; median R² = 0.711; no process truth",
    ]
    y = 232
    for line in lines:
        d.text((1080, y), line, font=font(12), fill="#29465d")
        y += 21

    d.rounded_rectangle((1055, 130, 1378, 181), radius=16, fill="#dff3e9", outline="#a9d7bc", width=2)
    d.text((1103, 139), "3 Evidence-tiered outputs", font=font(17, True), fill="#123f63")

    d.rounded_rectangle((1048, 348, 1380, 454), radius=12, fill="#d9f0df", outline="#7db892", width=2)
    d.text((1080, 371), "Lower Anayari", font=font(19, True), fill="#123f63")
    d.text((1080, 400), "Exchange/weathering hypotheses", font=font(13), fill="#4d9870")
    d.text((1280, 420), "PSI", font=font(12, True), fill="#4d9870")

    d.rounded_rectangle((1048, 448, 1380, 580), radius=12, fill="#ffe2d8", outline="#e68e72", width=2)
    d.text((1080, 467), "Talensi", font=font(19, True), fill="#123f63")
    d.text((1080, 496), "Mining/agriculture redox hypotheses", font=font(13), fill="#d36145")
    d.text((1280, 516), "PSI", font=font(12, True), fill="#d36145")

    d.rounded_rectangle((1058, 584, 1370, 651), radius=12, fill="#123f63", outline="#123f63", width=2)
    d.text((1090, 597), "Evidence-tiered diagnosis", font=font(18, True), fill="#ffffff")
    d.text((1098, 624), "transparent process hypotheses", font=font(14), fill="#d9edf7")

    # Correct the bottom validation-tier label.
    d.rounded_rectangle((820, 718, 1050, 760), radius=13, fill="#fff1e8", outline="#e8a07d", width=2)
    d.text((832, 729), "PHREEQC-compatible", font=font(12, True), fill="#123f63")
    d.text((842, 745), "feasibility proxy", font=font(12, True), fill="#123f63")
    out = io.BytesIO()
    im.save(out, format="PNG", optimize=True)
    return out.getvalue()


def update_supplementary() -> None:
    doc = Document(READY / "SupplementaryInformation_CG.docx")
    replace_paragraphs(
        doc,
        {
            46: "Supplementary Figure S1. Public USGS residence-time parity check using the M3 age-fraction-constrained benchmark. Among 1,272 rows, 1,249 had finite log10 errors; R² = 0.71, median absolute log10 error = 0.17, 61.4% were within a factor of two and 85.9% within a factor of ten. The USGS ages and age fractions are reported model outputs, so this is screening-level parity rather than direct true-age validation or full TracerLPM equivalence.",
            48: "Supplementary Figure S2. PHREEQC-compatible forward-feasibility proxy. Panel (a) shows the RMSE distribution (median = 0.062 mmol/L); panel (b) relates RMSE to Nash–Sutcliffe efficiency (median NSE = 0.999); and panel (c) shows percent bias (median PBIAS = 0.9%). The simulator uses locked PHREEQC-consistent saturation fields; live PHREEQC kinetic execution remains pending.",
            50: "Supplementary Figure S3. Field-scale chemical discovery fit for the Lower Anayari and Talensi groundwater networks. The histograms show edge-level chemistry R² for generated coordinate/head-proxy graphs: median R² = 0.60 for Lower Anayari and 0.77 for Talensi. These fits describe hydrochemical closure of a field demonstration and do not provide independent process-truth validation."
        },
    )
    set_table(
        doc.tables[0],
        [
            ["Data type", "Required field", "Unit", "Required/optional", "QC rule", "Used in module"],
            ["Location", "latitude/longitude or x/y", "decimal degrees or projected", "required", "valid coordinate range", "graph"],
            ["Elevation/head", "elevation, head_m", "m", "recommended", "non-missing preferred", "topology"],
            ["Major ions", "Ca, Mg, Na, K, HCO3, Cl, SO4, NO3", "mmol/L or mg/L", "required", "charge-balance screening", "chemistry"],
            ["Conditional ions", "F, Fe, PO4, Br", "mg/L or ug/L", "optional", "detection-limit handling", "logic gates"],
            ["Stable isotopes", "d18O, d2H", "permil", "optional/enhanced", "tracer range check", "transport"],
            ["Nuclear tracers", "tritium, 14C", "TU, pmC", "optional", "tracer range check", "age"],
            ["Field parameters", "pH, EC, temperature, DO/Eh", "field units", "recommended", "field range check", "geochemistry"],
        ],
    )
    set_table(
        doc.tables[1],
        [
            ["Process family", "Reaction/process term", "Chemical signature", "PHREEQC/SI rule", "Logic gate", "Allowed direction"],
            ["carbonate dissolution", "calcite", "Ca, HCO3", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["carbonate dissolution", "dolomite", "Ca, HCO3, Mg", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["evaporite dissolution", "gypsum", "Ca, SO4", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["evaporite dissolution", "halite", "Cl, Na", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["silicate weathering", "albite", "HCO3, Na", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["redox process", "pyrite_oxidation_aerobic", "Fe, SO4", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["nitrate input", "NO3src", "NO3", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["redox process", "denit", "HCO3, NO3", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["ion exchange", "CaNa_exch", "Ca, Na", "SI/forward check when available", "reaction dictionary", "signed"],
            ["ion exchange", "NaCa_exch", "Ca, Na", "SI/forward check when available", "reaction dictionary", "nonnegative"],
            ["ion exchange", "MgNa_exch", "Mg, Na", "SI/forward check when available", "reaction dictionary", "signed"],
            ["ion exchange", "NaMg_exch", "Mg, Na", "SI/forward check when available", "reaction dictionary", "nonnegative"],
        ],
    )
    set_table(
        doc.tables[2],
        [
            ["Parameter", "Value/range", "Used in module", "Scientific rationale"],
            ["n_realisations", "100", "benchmark", "Monte Carlo recovery ensemble"],
            ["lambda_l1 (default)", "0.002", "sparse inversion", "default sparse-recovery penalty"],
            ["AICc-selected lambda", "0.0483", "site model selection", "selected for Lower Anayari and Talensi only"],
            ["major_ion_rel_sigma", "0.04", "synthetic data", "analytical uncertainty on chemistry"],
            ["isotope_sigma_permil", "0.50", "isotope validation", "stable-isotope measurement noise"],
            ["transport_models_enabled", "evap, mix", "edge fitting", "minimum transport alternatives for M2"],
            ["PSI trials", "30 default", "field uncertainty", "edge stability under perturbation"],
            ["SI threshold", "0.2 diagnostic band", "forward checks", "thermodynamic feasibility screening"],
        ],
    )
    set_table(
        doc.tables[3],
        [
            ["Site", "Sample ID", "Latitude", "Longitude", "Elevation (m)", "Aquifer setting", "Land-use/mining influence", "Available tracers", "Completeness score"],
            ["Talensi", "TGW1", "10.71", "0.82", "217", "Crystalline basement", "Mining/Agri", "stable isotopes; no nuclear tracers", "1.00"],
            ["Talensi", "TGW2", "10.71", "0.81", "222", "Crystalline basement", "Mining/Agri", "stable isotopes; no nuclear tracers", "1.00"],
            ["Lower Anayari", "PUB1", "10.91", "-1.06", "217", "Alluvial/Basement", "Agriculture", "stable isotopes; no nuclear tracers", "1.00"],
            ["Lower Anayari", "PUB2", "10.91", "-1.06", "217", "Alluvial/Basement", "Agriculture", "stable isotopes; no nuclear tracers", "1.00"],
        ],
    )
    set_table(
        doc.tables[4],
        [
            ["Edge ID", "From node", "To node", "Distance (km)", "Elevation/head relation", "Edge confidence", "Age consistency", "Chemical match R2", "Dominant reaction", "Status"],
            ["Talensi_48_61", "TGW49", "THDW10", "2.52", "downgradient", "1.00", "field tracer absent", "1.00", "carbonate dissolution", "screening-level demonstration"],
            ["Talensi_45_44", "TGW46", "TGW45", "1.93", "upgradient/flat", "1.00", "field tracer absent", "1.00", "nitrate input", "screening-level demonstration"],
            ["Talensi_48_42", "TGW49", "TGW43", "2.54", "downgradient", "1.00", "field tracer absent", "1.00", "redox process", "screening-level demonstration"],
            ["Manu_30_28", "GOB38", "GOB36", "1.31", "downgradient", "1.00", "field tracer absent", "0.99", "evaporite dissolution", "screening-level demonstration"],
            ["Talensi_51_26", "TGW52", "TGW27", "3.31", "upgradient/flat", "1.00", "field tracer absent", "0.99", "nitrate input", "screening-level demonstration"],
            ["Talensi_28_27", "TGW29", "TGW28", "1.57", "downgradient", "1.00", "field tracer absent", "0.99", "nitrate input", "screening-level demonstration"],
        ],
    )
    for idx, name in enumerate([
        "Manuscript_Supp_FigS1_Public_Age_Validation.png",
        "Manuscript_Supp_FigS2_Geochemical_Validation.png",
        "Manuscript_Supp_FigS3_Ghana_Field_Residuals.png",
    ]):
        replace_image(doc, idx, image_bytes(FIG_DIR / name), f"Supplementary Figure S{idx+1}")
    doc.save(READY / "SupplementaryInformation_CG_updated.docx")


def update_highlights() -> None:
    doc = Document(READY / "Highlights.docx")
    replace_paragraphs(
        doc,
        {
            1: "Hydrosheaf integrates residence-time inversion, flow-topology inference, and inverse hydrogeochemistry in one auditable workflow.",
            2: "Evidence-constrained directed edges combine hydraulic, isotopic, chemical, and age-order information.",
            3: "Sparse L1 fitting ranks process classes and makes reaction identifiability limits explicit.",
            4: "Synthetic benchmarking gives chemistry R2 = 0.955, no-prior topology F1 = 0.618, and synthetic age R2 = 0.98.",
            5: "Public-age parity, PHREEQC proxy, and Ghana field results are reported as screening-level evidence with explicit caveats.",
        },
    )
    doc.save(READY / "Highlights_updated.docx")


def update_cover_letter() -> None:
    doc = Document(READY / "Cover_Letter- CG.docx")
    replace_paragraphs(
        doc,
        {
            11: "Evaluation is reported across synthetic recovery, MODFLOW/MODPATH topology comparison, synthetic and public residence-time benchmarking, a PHREEQC-compatible forward-feasibility proxy, and a Northern Ghana field-hydrochemistry demonstration. These tiers are deliberately separated because they test different claims; the PHREEQC tier remains a proxy and the Ghana networks have no independent process-truth labels.",
            12: "The results show chemistry R2 = 0.955 for the complete synthetic benchmark, no-prior topology F1 = 0.618 (precision 0.487, recall 0.845), synthetic network-age R2 = 0.98 with MAE = 183.3 years, public-age parity R2 = 0.71 with 61.4% within a factor of two, and 89.1% feasibility in the PHREEQC-compatible proxy. The Ghana demonstration generated interpretable process hypotheses but is explicitly presented as screening-level evidence rather than independent process validation.",
        },
    )
    doc.save(READY / "Cover_Letter- CG_updated.docx")


def update_graphical_abstract() -> None:
    doc = Document(READY / "Graphical abstract-CG.docx")
    replace_image(doc, 0, make_updated_graphical_abstract(), "Updated Hydrosheaf evidence-tiered graphical abstract")
    doc.save(READY / "Graphical abstract-CG_updated.docx")


def main() -> None:
    update_supplementary()
    update_highlights()
    update_cover_letter()
    update_graphical_abstract()
    print("Updated supplementary information, highlights, cover letter, and graphical abstract.")


if __name__ == "__main__":
    main()
