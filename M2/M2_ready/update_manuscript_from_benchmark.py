from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[2]
INPUT = ROOT / "M2" / "M2_ready" / "Manuscript-CG.docx"
OUTPUT = ROOT / "M2" / "M2_ready" / "Manuscript-CG_updated.docx"
FIG_DIR = ROOT / "M2" / "m2_benchmark" / "figures" / "Manuscript_Ready"


def replace_paragraph_text(document: Document, replacements: dict[int, str]) -> None:
    paragraphs = document.paragraphs
    for index, text in replacements.items():
        paragraph = paragraphs[index]
        # Preserve paragraph-level style and geometry; replace only the runs.
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)


def replace_embedded_image(document: Document, shape_index: int, image_path: Path) -> None:
    shape = document.inline_shapes[shape_index]
    rid = shape._inline.graphic.graphicData.pic.blipFill.blip.get(qn("r:embed"))
    rel = document.part.rels[rid]
    rel.target_part._blob = image_path.read_bytes()
    # Add a concise alternative description for screen readers and QA tools.
    shape._inline.docPr.set("descr", image_path.stem.replace("_", " "))
    shape._inline.docPr.set("title", image_path.stem.replace("_", " "))


def set_table(table, rows: list[list[str]]) -> None:
    # Keep the existing table style and geometry, but make the evidence table
    # match the current benchmark outputs.
    while len(table.rows) < len(rows):
        table.add_row()
    for r, row in enumerate(rows):
        for c, value in enumerate(row):
            if c < len(table.columns):
                table.cell(r, c).text = value
    # Blank any trailing cells/rows so stale values cannot survive.
    for r in range(len(rows), len(table.rows)):
        for c in range(len(table.columns)):
            table.cell(r, c).text = ""


def main() -> None:
    document = Document(INPUT)

    replacements = {
        30: (
            "Groundwater quality in data-limited aquifers is governed by the interplay of residence time, flow-path connectivity, and geochemical reactions, yet existing tools address these dimensions separately. This study introduces Hydrosheaf, a Python-based framework that couples residence-time inversion, directed graph-topology inference and sparse inverse hydrogeochemical modelling in one workflow. Sheaf-style consistency refinement scores candidate edges against hydraulic, isotopic and chemical evidence; network-enhanced Bayesian inference propagates graph constraints into age uncertainty; and L1-regularised reaction fitting uses PHREEQC-compatible logic gates to rank chemically feasible solutions. Across 100 locked synthetic realisations, median transport-parameter absolute error was 0.056 and median reaction-extent error was 0.050 mmol/L; active-reaction recovery in the refreshed figure was R2 = 0.57 (MAE = 0.33 mmol/L), with 32.7% of inactive terms exceeding the 0.05 mmol/L leakage threshold. No-prior topology recovery against MODPATH gave F1 = 0.618 (precision 0.487, recall 0.845), whereas the prior-assisted mode reproduced the reference edges exactly, as expected for a conditioning-prior test. Synthetic network-age recovery gave log-space R2 = 0.98 and MAE = 183.3 years. The independent public-age comparison was screening-level parity with reported USGS model outputs (n = 1,249 finite log errors; R2 = 0.71; median absolute log10 error = 0.17; 61.4% within a factor of two), not true-age validation. A deterministic PHREEQC-compatible mass-balance proxy marked 89.1% of edges thermodynamically feasible; live PHREEQC execution remains pending. The Northern Ghana demonstration generated 208 edges with overall chemistry R2 = 0.711 (Lower Anayari median 0.60; Talensi median 0.77) but no independent process-truth labels. Hydrosheaf therefore provides an auditable, evidence-tiered workflow for process diagnosis where calibrated numerical models are unavailable, while retaining explicit identifiability and transferability limits."
        ),
        44: (
            "Four methodological contributions are advanced. First, sheaf-style consistency refinement combines hydraulic, isotopic and chemical evidence to score directed-edge plausibility. Second, network-enhanced Bayesian age inference uses graph structure to narrow posterior uncertainty and evaluate age-order consistency; the resulting uncertainty reduction is a model property and should not be interpreted as a universal improvement in reference-age accuracy. Third, sparse L1-regularised reaction fitting with thermodynamic logic gates and geological-bias penalties ranks parsimonious candidate process classes while retaining non-uniqueness where the ion data are insufficient. Fourth, a process-stability index (PSI) quantifies how often reaction terms recur across Monte Carlo perturbations, distinguishing stable signals from data-contingent attributions."
        ),
        45: (
            "The evaluation comprises five evidence tiers: locked synthetic recovery with known truth; MODFLOW/MODPATH-referenced directed-topology comparison; synthetic and public residence-time benchmarking; PHREEQC-compatible forward-feasibility proxy testing; and a Northern Ghana field-hydrochemistry demonstration. These tiers are intentionally reported separately because they test different claims and do not constitute a single proof of universal field validity."
        ),
        111: (
            "A controlled synthetic benchmark evaluated recovery of known transport parameters and reaction extents from noisy chemistry data. The benchmark used a ground-truth YAML configuration specifying source chemistry, stable-isotope values, a directed generation network, and true evaporation, mixing and reaction parameters. Sequentially generated downstream vectors were perturbed with multiplicative Gaussian noise. One hundred stochastic realisations were generated per experiment, with complete-ion, ion-incomplete and head-absent scenarios and full, sparse and heuristic-only topologies. Performance was quantified using the displayed transport and reaction recovery diagnostics, chemistry R2, absolute error and noise-aware isotope metrics (Figure 3; Table 2)."
        ),
        113: (
            "The topology-validation benchmark evaluated Hydrosheaf against directed connectivity extracted from MODPATH particle endpoints and pathlines. Three operational modes were separated: prior-assisted conditioning, no-prior inference from spatial/elevational/head-proxy evidence, and threshold sensitivity. True-positive, false-positive and false-negative edges were tabulated to compute precision, recall and F1. This separation prevents circular interpretation of MODPATH priors as independent evidence and restricts the claim to topology conversion and recovery."
        ),
        117: (
            "A complementary metric suite was used across the five evidence tiers. Precision, recall and F1 quantified topology; chemistry R2, RMSE and NSE quantified synthetic and forward-proxy fit; AICc assessed regularisation parsimony; Kendall tau assessed downstream age ordering; and PSI quantified process recurrence across Monte Carlo perturbations. Public-age parity was reported in log10-error and within-factor metrics because the USGS targets are model-derived ages and age fractions rather than direct true ages."
        ),
        127: (
            "The nuclear-tracer module produced single-node MRT estimates for benchmark nodes with tracer data and network-enhanced Bayesian MRTs for nodes with non-zero in-degree. Network posterior interval widths were 52% narrower than the single-node intervals in this benchmark, by construction of the conjugate-updating model; this is an uncertainty-partitioning result, not an independent demonstration of lower reference-age error."
        ),
        131: (
            "Evaporation and mixing were recovered with different reliability across the 100 benchmark realisations. Mixing and dilution parameters were recovered whenever the relevant endmember was available, whereas evaporation-model classification was less reliable for the low-concentration-factor edge. Across all 600 transport rows, 553 had finite parameter estimates, with median absolute error 0.056 and an overall model-classification rate of 92.2%; evaporation classification was the principal source of misses. The positive bias in concentration-factor recovery is consistent with confounding between evaporative concentration and reaction-derived salinity, reinforcing the need for conservative-isotope evidence."
        ),
        132: (
            "The refreshed reaction-recovery panel contained 2,100 active and 4,900 inactive truth rows. Active rows showed R2 = 0.57 and MAE = 0.33 mmol/L, while 32.7% of inactive rows exceeded 0.05 mmol/L. The leakage was concentrated in chemically compensating terms, especially inactive albite and CaNa-exchange terms, and therefore represents sparse-recovery ambiguity rather than confident mineral identification. Reaction extents should consequently be reported as identifiability classes or ranked candidate processes, not unique phase histories."
        ),
        133: (
            "Performance under degraded data conditions remained informative but was not invariant. Median chemistry R2 was 0.955 for complete ions, 0.888 when major ions were withheld, and 0.955 when head data were absent. Missing major ions therefore weakened reaction discrimination, whereas head-data absence primarily affected graph construction. These results support graceful degradation, not equivalence of incomplete and complete datasets."
        ),
        135: (
            "The isotopic consistency check yielded pointwise R2 = 0.52 for delta18O, while the edge-mean comparison reached R2 = 0.99; the pointwise MAE was 0.58 per mil and the expected edge-difference noise was 0.71 per mil (Figure 3, Panel D). The defensible interpretation is noise-aware edge-mean consistency, not point-level isotope proof."
        ),
        138: (
            "Network Bayesian MRT estimates showed strong agreement with the locked synthetic ages in log space (R2 = 0.98; MAE = 183.3 years; Figure 5, Panel A). Posterior interval widths were 52% narrower than single-node intervals, reflecting the imposed conjugate-updating structure rather than an independently measured reduction in age error. The age-parity plot separated young, intermediate, old, fossil and mixed classes, with the greatest spread in intermediate and mixed classes where the age constraints are least informative."
        ),
        139: (
            "Posterior age distributions separated the young, intermediate, old, fossil and mixed synthetic classes, but the broad tails in the intermediate, mixed and fossil groups show that class separation does not imply precise point-age recovery. These distributions are therefore useful for screening and uncertainty communication rather than for asserting exact groundwater ages."
        ),
        142: (
            "The public USGS comparison provided the strongest external residence-time check currently available, but it was a parity assessment against reported model outputs and reported age fractions. Among 1,272 rows, 1,249 had finite log10 errors; R2 was 0.71, median absolute log10 error was 0.17, 61.4% were within a factor of two and 85.9% within a factor of ten (Figure 5, Panel B; Supplementary Figure S1). These results support screening-level agreement, not direct true-age validation or full TracerLPM equivalence."
        ),
        146: (
            "Comparison with the 174 MODPATH reference edges showed exact agreement in the prior-assisted mode: 174 true positives, zero false positives and zero false negatives (precision = recall = F1 = 1.00). This result demonstrates faithful transmission of a supplied physics prior; it is not independent topology inference."
        ),
        150: (
            "The regularisation path showed that the AICc minimum occurred at lambda = 0.0483 for both Lower Anayari and Talensi (Figure 6). The selected value is a site-specific balance between residual norm and model complexity; it should not be presented as a universal optimum."
        ),
        151: (
            "At lambda = 0.0483, residual norms remained substantially below the over-regularised solutions while AICc favoured a parsimonious model. This supports using AICc as a transparent selection rule for the analysed sites, while retaining the possibility that different aquifers, ion panels or reaction dictionaries will require different regularisation strengths."
        ),
        157: (
            "Forward feasibility was evaluated with a deterministic PHREEQC-compatible mass-balance proxy using the locked saturation fields. The proxy classified 89.1% of 1,000 synthetic edge realisations as thermodynamically feasible; median RMSE was 0.062 mmol/L and median NSE was 0.999. This is a screening-level feasibility result, not a live PHREEQC kinetic validation; an executable PHREEQC backend remains to be configured and rerun."
        ),
        158: (
            "Proxy residuals were heterogeneous across edges, and the QC output identifies cases with larger errors for follow-up. Because the current simulator is a deterministic mass-balance proxy rather than a live kinetic PHREEQC run, these residuals should guide additional field and thermodynamic checks rather than be interpreted as predictive performance."
        ),
        160: (
            "The framework was deployed on two contrasting Northern Ghana datasets: Lower Anayari, a crystalline-basement system, and Talensi, influenced by artisanal mining and agricultural nitrogen inputs. The datasets had sparse monitoring, no nuclear-tracer records and no calibrated regional flow model. The missing-data handler therefore configured a chemistry-and-coordinate workflow with optional tracer and thermodynamic modules inactive."
        ),
        161: (
            "Hydrosheaf generated process-discovery networks for both sites (Figure 4). Lower Anayari contained 82 generated edges, of which 66 selected mixing and 16 selected evaporation; Talensi contained 126 edges, of which 96 selected mixing and 30 selected evaporation. Median chemistry R2 was 0.60 for Lower Anayari, 0.77 for Talensi and 0.711 overall. These are mass-balance fits to observed downstream chemistry from a generated coordinate/head-proxy graph, not independent process-truth validation."
        ),
        162: (
            "High-confidence pathways were associated with cation exchange, nitrate cycling and lithology-dependent dissolution terms, but the relative ranking varied by edge. Lower Anayari showed frequent exchange and denitrification terms, whereas Talensi showed stronger nitrate-source and redox-related signatures. The process labels are therefore best treated as screening-level hypotheses to be corroborated with tracer, mineralogical and field-redox evidence."
        ),
        164: (
            "Field chemistry-fit distributions were heterogeneous, with median R2 = 0.60 for Lower Anayari and 0.77 for Talensi. Lower-fit edges occurred near the confidence boundary or where charge-balance and missing-indicator limitations reduced identifiability. The residual distribution is consequently a diagnostic for prioritising field verification, not a measure of independent process accuracy."
        ),
        168: (
            "Aerobic pyrite oxidation recurred in Talensi (PSI = 0.49) but was absent in Lower Anayari (PSI = 0.00), while gypsum dissolution was absent in Talensi (PSI = 0.00). These contrasts are consistent with the proposed mining/agriculture versus basement-weathering distinction, but they remain stability patterns under the analysed perturbations rather than guarantees of unique mineral sources."
        ),
        172: (
            "The central methodological contribution of Hydrosheaf is the integration of three normally isolated interpretive dimensions within a single computationally consistent workflow. Conventional approaches treat these separately: hydrogeochemists apply PHREEQC or NETPATH along user-specified paths (Parkhurst & Appelo, 2013; Roy et al., 2020); isotope hydrologists fit LPM models independently to individual samples (Ma et al., 2019; Casillas-Trasvina et al., 2022); and numerical modellers run MODFLOW/MODPATH to generate flow topology as a post-hoc interpretation layer (Juckem & Starn, 2021; Meyer et al., 2018). Hydrosheaf addresses this separation by making topology, residence-time uncertainty and reaction non-uniqueness visible in the same edge-level record. The integration constrains candidate interpretations and improves auditability, but it does not remove equifinality or replace calibrated numerical and geochemical models where those are available."
        ),
        180: (
            "MODFLOW/MODPATH provides physically rigorous particle-tracking topology, and its outputs can serve as validation references or optional priors in Hydrosheaf. In the Savage benchmark, the prior-assisted mode reproduced the supplied edge set exactly, whereas the independent no-prior mode achieved F1 = 0.618 (recall 0.845, precision 0.487). The latter demonstrates high recall with substantial overconnection for this case and supports screening-level use where a calibrated model is unavailable; it does not establish generalisation to heterogeneous three-dimensional field systems."
        ),
        187: (
            "The synthetic benchmark establishes computational recovery under controlled conditions: chemistry R2 was 0.955 for the complete-ion baseline, transport median absolute error was 0.056, and the refreshed reaction panel showed moderate active-row agreement (R2 = 0.57) with measurable inactive-term leakage. The results support mass-balance-consistent process classes and graceful degradation, but they cannot replace field-scale validation or prove unique reaction histories."
        ),
        188: (
            "Synthetic network-age recovery (log-space R2 = 0.98; MAE = 183.3 years) demonstrates consistency with the locked age design, while the 52% narrower posterior intervals are imposed by the network-updating model. The public USGS comparison is weaker but more independent: R2 = 0.71, median absolute log10 error = 0.17 and 61.4% within a factor of two. Because the public targets are reported model ages and fractions, this tier supports screening-level parity rather than true-age accuracy or full TracerLPM equivalence."
        ),
        189: (
            "The F1-score of 1.00 under prior-assisted conditions confirms the fidelity of MODPATH-prior ingestion but is not independent inference. The no-prior F1-score of 0.618 (recall 0.845, precision 0.487) demonstrates high-recall heuristic recovery with substantial overconnection in the tested benchmark. These results support transparent prior/no-prior separation and edge-level corroboration before geochemical interpretation."
        ),
        190: (
            "The PHREEQC-compatible mass-balance proxy marked 89.1% of synthetic edges feasible, with median RMSE = 0.062 mmol/L and median NSE = 0.999. The proxy supports thermodynamic screening and mass-balance consistency; it does not substitute for a live PHREEQC kinetic forward run, which remains pending."
        ),
        193: (
            "Hydrosheaf is designed for settings in which a calibrated regional flow model is absent. The Ghana demonstration shows that chemistry-only process-discovery networks can be generated from sparse field inputs, but their interpretation must remain screening-level because the graph is generated from coordinates and head proxies and lacks independent process-truth labels. PSI-ranked edges can support conceptual-model development, monitoring prioritisation and targeted follow-up sampling, provided that high-impact decisions are not based on a single inferred reaction term."
        ),
        195: (
            "Talensi showed a recurring pyrite-oxidation signal (PSI = 0.49) and nitrate cycling, while Lower Anayari showed stronger exchange signatures and no pyrite oxidation. These patterns are useful hypotheses for water-quality risk assessment and sampling design, but the PSI values express recurrence under the analysed perturbations, not the prevalence or uniqueness of the underlying mineral source."
        ),
        198: (
            "Input completeness directly governs output reliability. Missing indicator ions, charge-balance error and absent tracers widen the set of chemically equivalent solutions. Graph assumptions—search radius, confidence threshold and elevation/head proxies—can also generate overconnected or missed edges when topography is decoupled from the potentiometric surface."
        ),
        199: (
            "L1 regularisation and thermodynamic constraints do not eliminate non-uniqueness under sparse data. The refreshed reaction panel quantified this limitation through moderate active-row agreement and false-positive inactive terms. Low PSI communicates instability but does not identify the true process; high-PSI reactions remain screening-level attributions requiring independent geochemical, mineralogical or tracer corroboration before management use."
        ),
        200: (
            "Each validation tier has transferability limits. The synthetic benchmark cannot represent all aquifer chemistries; the public-age comparison is parity to reported USGS model outputs; the MODPATH test uses one reference model; the PHREEQC evidence is a proxy; and the Ghana demonstration covers two sites without process-truth labels. Claims should therefore be restricted to the tested configurations and extended only after site-specific validation."
        ),
        208: (
            "Hydrosheaf demonstrates that residence-time inversion, directed graph-topology inference and sparse inverse hydrogeochemical modelling can be combined in one auditable workflow for groundwater process diagnosis in data-limited aquifers. Its principal contribution is the replacement of manually assumed flow paths with probabilistically scored edges that can be constrained by hydraulic, isotopic, chemical and particle-tracking evidence. Thermodynamic logic gates, geological-bias penalties and L1 regularisation rank parsimonious process classes while retaining explicit uncertainty and non-uniqueness."
        ),
        209: (
            "The evidence tiers were not equivalent. Synthetic benchmarking gave chemistry R2 = 0.955, transport median absolute error = 0.056 and reaction median error = 0.050 mmol/L; the refreshed reaction panel showed active-row R2 = 0.57, MAE = 0.33 mmol/L and 32.7% inactive-term leakage. MODPATH comparison gave F1 = 1.00 only when the reference prior was supplied and F1 = 0.618 without that prior. Synthetic network-age recovery gave R2 = 0.98 and MAE = 183.3 years. Public-age parity was R2 = 0.71 with 61.4% within a factor of two, while the PHREEQC-compatible proxy marked 89.1% of edges feasible. These results support a tiered, screening-level reliability claim, not universal process validation."
        ),
        210: (
            "The Ghana demonstration generated 208 process-discovery edges from sparse chemistry and coordinate/head-proxy inputs. Median chemistry R2 was 0.60 for Lower Anayari and 0.77 for Talensi. Exchange and nitrate-cycling terms were common at both sites, while pyrite oxidation recurred in Talensi and was absent in Lower Anayari. These distinctions are useful for hypothesis generation, but the absence of an independent field process-truth graph means that they require tracer, mineralogical and redox corroboration."
        ),
        211: (
            "Overall, Hydrosheaf provides an open and reproducible framework for evidence-tiered groundwater process diagnosis in sub-Saharan Africa and other data-limited regions. Its contribution is strongest as an auditable integration and prioritisation layer that makes topology, residence-time uncertainty and geochemical non-uniqueness explicit."
        ),
        331: "Figure 2. Topology-only comparison of Hydrosheaf-inferred directed connectivity against the MODPATH reference simulation. The prior-assisted mode tests faithful physics-prior ingestion; the no-prior mode tests independent heuristic recovery and is the relevant capability metric.",
        334: "Figure 3. Synthetic benchmark evidence across 100 locked realisations. (A) physical transport-parameter recovery; (B) reaction extents for active truth rows and inactive-term sparsity, with active-row R2 = 0.57, MAE = 0.33 mmol/L and 32.7% of inactive terms exceeding 0.05 mmol/L; (C) chemistry-fit robustness under missing major ions and absent head data; and (D) noise-aware isotope recovery, with point R2 = 0.52, edge-mean R2 = 0.99, point MAE = 0.58 per mil and edge-difference noise sigma = 0.71 per mil. These are benchmark evidence, not proof of unique process recovery.",
        338: "Figure 4. Data-limited process-discovery networks for Lower Anayari and Talensi. The graphs are generated from field coordinates and head proxies; edge colours denote inferred process families and edge annotations show PSI for selected links. The figure is a field-hydrochemistry demonstration without an independent process-truth graph, tracer-age validation or MODPATH field benchmark.",
        341: "Figure 5. Residence-time benchmarking. (a) Synthetic network Bayesian recovery gives log-space R2 = 0.98 and MAE = 183.3 years. (b) The public USGS comparison is age-fraction-constrained parity to reported model outputs (1,249 finite log errors), with R2 = 0.71, median absolute log10 error = 0.17, 61.4% within a factor of two and 85.9% within a factor of ten. Public reported ages are model-derived targets, not direct true ages or evidence of full TracerLPM equivalence.",
        343: "Figure 6. Site-specific AICc regularisation selection for Lower Anayari and Talensi. The minimum occurs at lambda = 0.0483 in both panels. This value is a transparent choice for the analysed datasets, not a universal optimum for other aquifers, ion panels or reaction dictionaries.",
        346: "Figure 7. Detailed process-stability matrix (PSI) for Lower Anayari and Talensi. Values express identification probability under the analysed Monte Carlo perturbations. The contrasts between cation exchange, nitrate cycling, gypsum and aerobic pyrite oxidation are stability patterns requiring independent corroboration; the legacy filename should not be interpreted as a guarantee."
    }
    replace_paragraph_text(document, replacements)

    # Main-text tables (the manuscript has several equation-only tables before
    # these; indices 15-20 are the visible evidence tables).
    set_table(
        document.tables[16],
        [
            ["Validation tier", "Dataset/source", "What is tested", "Reference/target", "Main metric"],
            ["Synthetic benchmark", "100 locked realisations", "transport, reaction and chemistry recovery", "known truth", "chemistry R2=0.955; transport median AE=0.056; reaction median AE=0.050"],
            ["Public-age parity", "USGS public-supply age release; M3 age-fraction parity", "reported age and age-fraction agreement", "USGS reported model outputs", "n=1272; R2=0.71; median |log10 error|=0.17; factor 2=61.4%"],
            ["MODPATH topology", "USGS Savage MODFLOW/MODPATH archive", "directed-edge recovery", "MODPATH edges", "prior-assisted F1=1.00; no-prior F1=0.618 (P=0.487, R=0.845)"],
            ["PHREEQC-compatible proxy", "locked saturation fields", "forward feasibility", "mass-balance proxy", "feasible=0.891; median RMSE=0.062 mmol/L; median NSE=0.999; live PHREEQC pending"],
            ["Ghana field demonstration", "Lower Anayari/Talensi", "generated-graph process discovery", "hydrochemical consistency only", "208 edges; overall median R2=0.711; no independent process truth"],
        ],
    )
    set_table(
        document.tables[18],
        [
            ["Validation group", "Reference age range (y)", "Hydrosheaf inferred range (y)", "R2", "MAE (y)", "Age-order consistency", "Interpretation"],
            ["Young/Recharge", "5–25", "3.7–34.4", "0.97", "0.81", "0.85", "synthetic age-class recovery"],
            ["Intermediate", "60–260", "38.5–433.3", "0.85", "15.75", "0.85", "synthetic age-class recovery"],
            ["Old/Deep", "520–2500", "305.6–7370.6", "0.82", "188.49", "0.85", "synthetic age-class recovery"],
            ["Overall synthetic", "5–4500", "3.7–9949.8", "0.98 (log-space)", "183.3", "0.85", "network Bayesian recovery"],
            ["Public USGS parity", "1–741000", "0–50000", "0.71 (log-space)", "median |log10| = 0.17", "not applicable", "screening-level parity to reported model outputs"],
        ],
    )
    set_table(
        document.tables[20],
        [
            ["Site", "Flow path/edge", "Dominant process", "Extent (mmol/L)", "Selected lambda", "Fit metric", "PSI", "Interpretation"],
            ["Talensi", "Talensi_48_61", "carbonate dissolution", "1.23", "0.0483", "objective 0.22 / R2 1.00", "1.00", "carbonates signal"],
            ["Talensi", "Talensi_45_44", "nitrate input", "2.21", "0.0483", "objective 0.82 / R2 1.00", "1.00", "nitrate hypothesis"],
            ["Talensi", "Talensi_48_42", "redox process", "4.12", "0.0483", "objective 0.41 / R2 1.00", "1.00", "redox hypothesis"],
            ["Lower Anayari", "Manu_30_28", "evaporite dissolution", "0.33", "0.0483", "objective 0.25 / R2 0.99", "1.00", "screening-level process label"],
        ],
    )

    figure_names = [
        "Manuscript_Fig1_Architecture.png",
        "Manuscript_Fig2_Topology_Validation.png",
        "Manuscript_Fig3_Synthetic_Validation.png",
        "Manuscript_Fig4_Ghana_Process_Network.png",
        "Manuscript_Fig5_Residence_Time_Validation.png",
        "Manuscript_Fig6_Optimal_Model_Selection.png",
        "Manuscript_Fig7_PSI_Robustness_Guarantee.png",
    ]
    for i, name in enumerate(figure_names):
        path = FIG_DIR / name
        if path.exists():
            replace_embedded_image(document, i, path)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
