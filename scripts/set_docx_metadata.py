"""Set required DOCX title metadata from each manuscript's authoritative H1."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
PAIRS = [
    (ROOT / "M5" / "manuscript" / "Manuscript-Final.md", ROOT / "M5" / "manuscript" / "Manuscript-Final.docx"),
    (ROOT / "M6" / "m6_field_transfer_benchmark" / "manuscript" / "Manuscript-Final.md", ROOT / "M6" / "m6_field_transfer_benchmark" / "manuscript" / "Manuscript-Final.docx"),
    (ROOT / "M7" / "m7_nonuniqueness_benchmark" / "manuscript" / "Manuscript-Final.md", ROOT / "M7" / "m7_nonuniqueness_benchmark" / "manuscript" / "Manuscript-Final.docx"),
    (ROOT / "M5" / "manuscript" / "supplementary" / "Supplementary-Figures-and-Tables.md", ROOT / "M5" / "manuscript" / "supplementary" / "Supplementary-Figures-and-Tables.docx"),
]


def main() -> int:
    for markdown, docx in PAIRS:
        title = next(line[2:].strip() for line in markdown.read_text(encoding="utf-8").splitlines() if line.startswith("# "))
        document = Document(docx)
        document.core_properties.title = title
        for section in document.sections:
            section.page_width = Inches(8.27)
            section.page_height = Inches(11.69)
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        for table in document.tables:
            table.autofit = True
            for row_index, row in enumerate(table.rows):
                tr_pr = row._tr.get_or_add_trPr()
                if row_index == 0:
                    header = OxmlElement("w:tblHeader")
                    header.set(qn("w:val"), "true")
                    tr_pr.append(header)
                else:
                    cant_split = OxmlElement("w:cantSplit")
                    tr_pr.append(cant_split)
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.paragraph_format.space_before = Pt(0)
                        paragraph.paragraph_format.space_after = Pt(0)
                        paragraph.paragraph_format.line_spacing = 1
                        for run in paragraph.runs:
                            run.font.size = Pt(6.5)
        document.save(docx)
        print(f"set title: {docx} -> {title}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
