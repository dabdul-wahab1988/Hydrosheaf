"""Remove legacy revision colour and inherited page-break sections from M3 DOCX files."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import RGBColor


REPO = Path(__file__).resolve().parents[3]
DOC_DIR = REPO / "M3" / "M3_geochemistry"


def remove_revision_red(doc: Document) -> int:
    changed = 0
    containers = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                containers.extend(cell.paragraphs)
    for paragraph in containers:
        for run in paragraph.runs:
            if str(run.font.color.rgb or "").upper() == "EE0000":
                run.font.color.rgb = RGBColor(0, 0, 0)
                changed += 1
    return changed


def main() -> None:
    main_path = DOC_DIR / "Manucript_3.docx"
    main_doc = Document(main_path)
    recoloured = remove_revision_red(main_doc)
    main_doc.save(main_path)

    supp_path = DOC_DIR / "Supplementary_Information_M3.docx"
    supp = Document(supp_path)
    matches = [p for p in supp.paragraphs if p.text.startswith((
        "For the N = 329 strict reportable subset",
        "For the strict reportable subset (N = 329)",
    ))]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one supplementary metric paragraph; found {len(matches)}")
    matches[0].text = (
        "For the strict reportable subset (N = 329), median absolute log₁₀ reference discrepancy was 0.0279, "
        "log₁₀ RMSE was 0.2769, and 87.5% were within a factor of two. The fraction sensitivity (N = 289) "
        "had median discrepancy 0.0216 and RMSE 0.1964. These quantify agreement with model-derived LPM "
        "outputs, not independently known true ages."
    )
    # Attach the first orientation change to the preceding content paragraph.
    # The legacy file stored it in a separate empty paragraph, which Word placed
    # alone on a blank portrait page before the landscape reproducibility table.
    paragraphs = list(supp.paragraphs)
    section_breaks = [p for p in paragraphs if p._p.xpath("./w:pPr/w:sectPr")]
    first_break = section_breaks[0]
    if not first_break.text.strip():
        first_index = next(i for i, p in enumerate(paragraphs) if p._p is first_break._p)
        target_index = first_index - 1
        while target_index >= 0 and not paragraphs[target_index].text.strip():
            target_index -= 1
        target = paragraphs[target_index]
        sect_pr = first_break._p.pPr.sectPr
        first_break._p.pPr.remove(sect_pr)
        target._p.get_or_add_pPr().append(sect_pr)
        for paragraph in paragraphs[target_index + 1:first_index + 1]:
            paragraph._element.getparent().remove(paragraph._element)

    supp.sections[1].start_type = WD_SECTION.NEW_PAGE
    for section in supp.sections[2:]:
        section.start_type = WD_SECTION.CONTINUOUS
    remove_revision_red(supp)
    supp.save(supp_path)
    print(f"Recoloured {recoloured} legacy-red manuscript runs; compacted supplementary sections.")


if __name__ == "__main__":
    main()
