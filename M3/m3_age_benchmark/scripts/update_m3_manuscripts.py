"""Synchronise M3 submission DOCX files with the locked tables and figures."""

from __future__ import annotations

import math
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO = Path(__file__).resolve().parents[3]
DOC_DIR = REPO / "M3" / "M3_geochemistry"
TABLE_DIR = REPO / "M3" / "m3_age_benchmark" / "tables" / "Manuscript_Ready"
FIG_DIR = REPO / "M3" / "m3_age_benchmark" / "figures" / "Manuscript_Ready"


MAIN_REPLACEMENTS = {
    "Groundwater age inference remains non-unique": (
        "Groundwater-age inference remains non-unique because tracer observations can be explained by alternative transit-time distributions, correction assumptions, tracer subsets, and spatial-connectivity hypotheses. We evaluated the residence-time and graph-prior component of Hydrosheaf using 1,272 harmonised rows from a public USGS groundwater-age release. Published USGS lumped-parameter-model (LPM) outputs were treated as model-derived reference products, not error-free true ages. An explicit identifiability and reportability guard retained 329 strict reported-configuration fits. On that reportable subset, strict configuration emulation had median absolute log₁₀ reference discrepancy 0.0279, log₁₀ RMSE 0.2769, log₁₀ R² 0.9371, and 87.5% within a factor of two. Adding reported age fractions yielded 289 reportable fits with median discrepancy 0.0216, RMSE 0.1964, R² 0.9698, and 91.7% within a factor of two, but this is a non-independent reported-output-constrained sensitivity because the fractions and reference ages share USGS LPM provenance. In the 329-node graph-age benchmark, weak hydraulic-proxy regularisation changed RMSE by −0.00135 log₁₀ while reducing within-factor-two agreement from 87.5% to 86.9%; therefore no candidate graph met the robust-improvement rule. Wrong-direction and randomised weak controls increased RMSE by 0.1008 and 0.6550 log₁₀, respectively. Leakage-guarded tracer withholding also showed no meaningful candidate-graph gain: ³H RMSE was 20.8181 TU for the single-well baseline and 20.8181–20.8415 TU for candidate graphs; SF₆ candidate graphs increased RMSE by 0.5–3.0%; and the small ¹⁴C change was negligible. CFC graph effects were non-estimable because the 4 CFC-11 and 6 CFC-12 reportable fits formed no eligible graph edges. Hydrosheaf M3 is therefore supported as a controlled diagnostic and topology-falsification benchmark, not as field validation, a source of true groundwater ages, or evidence of universal graph benefit."
    ),
    "The objectives are to:": (
        "The objectives are to compare reported-configuration LPM emulation, reported-output-constrained age-fraction sensitivity, Hydrosheaf-selected transit-time distributions, tracer and correction ablations, and graph-prior controls; quantify reference-workflow agreement only on identifiable reportable fits; test concentration prediction by target-tracer withholding; and determine whether graph topology provides defensible information beyond well-wise inference. A graph result is accepted only when RMSE improves, within-factor agreement does not deteriorate, and negative controls do not show a comparable effect."
    ),
    "Two further releases were obtained": (
        "Two further public releases were assessed for restricted use. The Western Principal Aquifers release overlaps the national release and lacks the coordinate coverage required for the graph benchmark, so it was not pooled into the active matrix. The Mississippi River Valley alluvial aquifer release provides useful site and tracer context, but the reported LPM table required for like-for-like model-derived reference ages was unavailable in the archived inputs. It was therefore excluded from parity, graph-age, and replication metrics. Neither release is presented as an independent replication of the reported-configuration benchmark."
    ),
    "Age classes were defined for diagnostic stratification": (
        "Age classes were defined for diagnostic stratification as modern (≤50 yr), intermediate (50–1,000 yr), old (1,000–30,000 yr), and very old (>30,000 yr). These are classes of model-derived reference output, not observed groundwater-age truth. The final audit found no supported paired raw and corrected atmospheric-equivalent gas values. Consequently, no empirical gas-correction effect is estimated, Supplementary Table S4 and Supplementary Figure S1 are withdrawn, and gas correction remains a documented modelling requirement rather than a validated result of this dataset."
    ),
    "The scenarios included reported-corrected parity": (
        "The active scenario matrix included reported-model parity, raw-gas and raw-¹⁴C ablations, screened young-gas correction, no-⁴He ablation, young-only and old-only tracer subsets, ¹⁴C ensemble handling, ⁴He uncertainty treatment, Hydrosheaf-selected TTDs, strict reported-configuration emulation, and reported-output-constrained age-fraction sensitivity. The hierarchical old-water prior was withdrawn because its pooled same-release priors were not independent and its archived diagnostic result was catastrophic (full-output RMSE 1.310 log₁₀ and R² 0.004). It is retained only for provenance and is not used in active tables, figures, rankings, or conclusions."
    ),
    "The scenario matrix was fixed before scenario-level interpretation": (
        "The active matrix used all 1,272 national-release rows and applied a reference-free reportability guard before any reference-age metric. Strict reported-configuration emulation produced 329 reportable fits, reported-output-constrained fraction sensitivity 289, Hydrosheaf model selection 309, screened young-gas correction 66, and reported-model parity 66. Direct mode comparisons used common support (N = 40); full-available rows are descriptive because supports differ."
    ),
    "Tracer-withholding cross-validation was used as a stronger predictive test": (
        "Tracer-withholding cross-validation tested prediction rather than agreement with model-derived ages. For every eligible row, the target was removed before model selection; the single-node fit and every graph-node age were then refitted without that target. Graph regularisation therefore received no full-fit neighbour ages containing the withheld tracer. Predictions were scored in native units only after the reference-free reportability guard. The withheld concentration was not used for fitting, graph construction, or hyperparameter tuning."
    ),
    "Computational reproducibility was controlled": (
        "Computational reproducibility was controlled through deterministic row selection, graph randomisation seed 42, a 90-step age grid, explicit scenario identifiers, and machine-readable manifests. The locked manifest records input and output SHA-256 hashes, parameters, reportability counts, graph-edge counts, software versions, and generated table and figure assets. Archived diagnostic scenarios are separated from active reportable outputs."
    ),
    "The benchmark used a national-scale public groundwater-age structure": (
        "The active benchmark used 1,272 harmonised rows from the national public-supply release. Reference-agreement results are restricted to identifiable reportable fits: N = 329 for strict reported-configuration emulation, N = 289 for the fraction sensitivity, and N = 309 for Hydrosheaf model selection. Tables 2–6 and Figures 2–7 use the same locked supports and distinguish reference discrepancy from target-withheld concentration prediction."
    ),
    "On the largest common-support comparison": (
        "On the N = 40 common support, strict reported-configuration emulation had median absolute log₁₀ reference discrepancy 0.0752, RMSE 0.3040, and 75.0% within a factor of two; the reported-output-constrained fraction sensitivity had 0.0653, 0.2854, and 75.0%, respectively. These differences describe reproduction of the published LPM workflow, not independent groundwater-age validation."
    ),
    "Full-available summaries confirmed the same broad direction": (
        "On their full reportable supports, strict reported-configuration emulation (N = 329) had median absolute log₁₀ reference discrepancy 0.0279, RMSE 0.2769, R² 0.9371, and 87.5% within a factor of two. The fraction sensitivity (N = 289) had 0.0216, 0.1964, 0.9698, and 91.7%, respectively. Hydrosheaf model selection (N = 309) had median discrepancy 0.1304, RMSE 0.6098, R² 0.7639, and 67.3% within a factor of two. Unequal supports and shared provenance preclude field-accuracy or superiority claims."
    ),
    "The graph-topology benchmark showed": (
        "The graph-age benchmark was evaluated only on the 329 strict reportable nodes. A lower RMSE was insufficient if within-factor-two agreement deteriorated. Supplementary Figure S2 reports edge geometry; Figure 4 and Table 4 report the locked graph effects."
    ),
    "Table 4 showed that the graph-age RMSE benchmark": (
        "Weak hydraulic-proxy regularisation produced the smallest candidate RMSE change (−0.00135 log₁₀), but within-factor-two agreement fell from 0.8754 to 0.8693. It therefore failed the robust-improvement rule. Weak wrong-direction and randomised controls increased RMSE by 0.1008 and 0.6550 log₁₀, respectively. No MRVA replication metric is reported because comparable reported LPM outputs were unavailable."
    ),
    "The wrong-direction controls were also informative": (
        "Wrong-direction and randomised controls demonstrate sensitivity to imposed topology, but their degradation is not positive evidence that a candidate graph is correct. Candidate graphs were neutral or harmful under the joint decision rule, while implausible controls were more harmful."
    ),
    "Tracer-withholding cross-validation showed no graph family improved tritium prediction": (
        "Leakage-guarded withholding retained 121 reportable ³H fits from 794 eligible rows and 75 reportable SF₆ fits from 262 eligible rows. For ³H, RMSE was 20.818104 TU at baseline, 20.818055 for the hydraulic graph, 20.841494 for the depth graph, and 21.566597 for the randomised control. For SF₆, RMSE was 2.843988 pptv at baseline and increased to 2.857993, 2.928528, and 2.928427 pptv. No young-tracer graph gain is supported."
    ),
    "Graph regularisation degraded or failed to improve inference": (
        "For ¹⁴C, 169 reportable fits remained from 1,103 eligible rows. RMSE was 26.830072 pmC at baseline, 26.829250 under the hydraulic graph, 26.902373 under the depth graph, and 27.819135 under the randomised control. The hydraulic difference (−0.0031%) is numerically negligible."
    ),
    "The result was therefore uniformly non-positive": (
        "Across reportable target-withheld tests, candidate graph effects were negligible or adverse. CFC-11 and CFC-12 retained only 4 and 6 reportable fits, respectively, with zero eligible graph edges; identical scenario RMSE values are a non-estimable graph comparison, not evidence of equivalence."
    ),
}

MAIN_REPLACEMENTS.update({
    "The study observed that modern and intermediate waters": (
        "On the strict reportable subset, median absolute log₁₀ reference discrepancy was 0.0890 for modern (N = 76), 0.0093 for intermediate (N = 85), 0.0233 for old (N = 158), and 0.0330 for very old water (N = 10). These strata describe agreement with model-derived reference outputs and have highly unequal support; they do not demonstrate age-class field accuracy."
    ),
    "Young-tracer discordance explained part": (
        "Young-tracer discordance is retained as an exploratory diagnostic of disagreement among available young-water proxies. It is not a causal explanation of reference discrepancy, and SF₆/CFC interpretation remains conditional on contamination, excess-air, recharge-temperature, atmospheric-history, and unsaturated-zone assumptions."
    ),
    "Old-water diagnostics were also heterogeneous": (
        "Old-water indicators were heterogeneous: 176 agreement, 3 conflict, 8 tension, 633 single-tracer, and 452 no-constraint cases. These are diagnostic classifications from available ¹⁴C and ⁴He products, not independent age validation or a basis for the withdrawn hierarchical prior."
    ),
    "The ablation results showed that reduced tracer information": (
        "Ablation effects were support-limited. On 65 paired reportable rows, removing ⁴He produced no change in median or mean log₁₀ discrepancy and no factor-of-two gains or losses. Several old-water scenarios likewise had near-zero paired effects on small supports. These results do not establish that a tracer is unimportant."
    ),
    "A direct SF6-removal experiment was not included": (
        "A direct SF₆-removal experiment was not included. The raw-gas scenario produced no reportable metrics, and the paired raw-versus-corrected atmospheric-equivalent audit had zero supported pairs. The gas-correction table and figure are withdrawn. Withheld SF₆ shows no graph gain, while CFC graph effects are non-estimable because no reportable CFC nodes are connected."
    ),
    "Carbon-14 simplification also required caution": (
        "Raw-¹⁴C and old-water ensemble perturbations had small, support-limited paired effects, and graph regularisation did not meaningfully improve withheld ¹⁴C. The young-only scenario retained 65 metric rows with RMSE 0.6872 and 53.8% within a factor of two, whereas the old-only scenario produced no reportable fits. This shows limited identifiability, not a universal tracer-suite rule."
    ),
    "The study observed five main results": (
        "Five results define the defensible claim. First, only 329 of 1,272 strict fits passed the reference-free reportability guard. Second, strict and age-fraction modes reproduced model-derived USGS outputs, but the fraction result is non-independent. Third, no candidate graph robustly improved the 329-node reference benchmark. Fourth, leakage-guarded ³H, SF₆, and ¹⁴C withholding showed no meaningful graph gain, and CFC graph comparisons were non-estimable. Fifth, the hierarchical old-water prior and unsupported gas-correction comparison were withdrawn. M3 is a controlled capability and falsification benchmark, not field validation."
    ),
    "The results show that this network enhancement is most valuable": (
        "The network component is most defensible as a falsification and sensitivity tool. Candidate graphs did not meet the joint improvement rule, while wrong-direction and randomised controls caused larger degradation. Control failure can reject an imposed topology; it cannot confirm that a surviving candidate graph is a true flow path."
    ),
    "Two potential misreadings": (
        "Three misreadings must be avoided. Agreement with published LPM outputs is not accuracy against true groundwater age; reported fractions and reference ages share provenance; and target-withheld tests are heavily filtered by reportability and sparse graph connectivity. The corrected tests support no meaningful graph gain and do not validate the complete Hydrosheaf framework."
    ),
    "The results show that directionality is particularly important": (
        "Wrong-direction controls were more harmful than the baseline, showing that the implementation responds to edge orientation. This is evidence that direction is consequential under the model, not evidence that the candidate direction represents actual groundwater flow."
    ),
    "The tracer-specific evidence in Figure 6": (
        "Figure 6 shows no meaningful candidate-graph improvement for ³H, SF₆, or ¹⁴C after every node was refitted without the target tracer. Randomised controls were more harmful for ³H and ¹⁴C, but essentially indistinguishable from the depth graph for SF₆. Graph value is not demonstrated for target-withheld prediction."
    ),
    "The study contributes to nuclear-isotope groundwater dating": (
        "The contribution to nuclear-isotope groundwater dating is methodological: target-tracer withholding, reference-free reportability checks, and negative controls expose unsupported topology effects. The corrected ³H benchmark retained 121 reportable fits and showed a negligible −0.0002% hydraulic-graph RMSE change and +0.112% depth-graph change."
    ),
    "The strongest broad result is not a graph-only result": (
        "The closest reference agreement was obtained when reported configurations or reported age fractions were used. Because those quantities reproduce the USGS LPM workflow, they establish implementation parity and sensitivity, not independent truth."
    ),
    "For data-limited aquifers, Hydrosheaf is most useful": (
        "For data-limited aquifers, the present evidence supports Hydrosheaf as a diagnostic workflow rather than a black-box age estimator. Reduced-tracer scenarios often became non-reportable or support-limited, so sampling recommendations remain hypotheses for independent field testing."
    ),
    "This modular positioning strengthens Hydrosheaf": (
        "This modular positioning prevents M3 from being misread as whole-framework validation. M3 evaluates only residence-time fitting, reference-workflow emulation, reportability, and graph-prior diagnostics. Other Hydrosheaf modules require separate validation."
    ),
    "The first risk is false connectivity": (
        "The principal risks are false connectivity, tracer contamination, non-identifiable fits, reference circularity, sparse reportable support, and transfer beyond the evaluated USGS release. Published LPM ages are model-derived targets; graph edges are hypotheses; and no Ghana-specific or semi-arid validation dataset was evaluated."
    ),
    "Compared with existing TracerLPM-style workflows": (
        "Strict reported-configuration mode is an emulation test by construction. Hydrosheaf model selection has larger reference discrepancy on its own reportable subset, but unequal supports prevent a general ranking. Hydrosheaf contributes reportability, target withholding, and topology-failure tests; it does not establish true ages or displace established LPM interpretation."
    ),
    "The benchmark shows that age-fraction constraints": (
        "Reported age fractions reduce reference discrepancy, but shared provenance prevents treating them as independent evidence. A defensible workflow harmonises inputs, establishes identifiability, separates emulation from prediction, tests target withholding, and rejects graph configurations that fail joint metric and negative-control checks."
    ),
    "Model validation should include single-well baselines": (
        "Future validation should include independent field constraints, reference-free reportability, target-tracer withholding with every graph node refit, graph-strength sensitivity, negative controls, and frozen manifests. A graph effect should replicate on independent aquifers before it is accepted."
    ),
    "The principal methodological contribution of this benchmark": (
        "The principal contribution is an auditable procedure for testing and rejecting topology assumptions. On 329 reportable nodes, no candidate graph robustly improved reference agreement: the smallest RMSE change was −0.00135 log₁₀ but within-factor-two agreement worsened. Wrong-direction and randomised weak controls increased RMSE by 0.1008 and 0.6550 log₁₀."
    ),
    "The independent predictive test was negative for tritium": (
        "Leakage-guarded tracer withholding found no meaningful candidate-graph gain. ³H RMSE was 20.818104 TU at baseline versus 20.818055 and 20.841494 for candidate graphs; SF₆ candidate graphs increased RMSE by 0.5–3.0%; and the smallest ¹⁴C change was −0.0031%. CFC graph comparisons were non-estimable because no eligible edges remained."
    ),
    "The strongest broad-sample configuration": (
        "Strict reported-configuration emulation reproduced model-derived outputs closely on 329 reportable fits, and the fraction sensitivity was closer on 289 fits. These are implementation-agreement and sensitivity results only. The 943 non-reportable strict rows, unequal supports, and shared provenance preclude population-wide or true-age claims."
    ),
    "The practical contribution is a defensible workflow": (
        "The practical contribution is a reproducible workflow recording when residence-time inference is reportable, when topology is harmful, and when an analysis must be withdrawn. The hierarchical old-water prior and unsupported gas-correction comparison remain only as archived failure evidence."
    ),
    "Hydrosheaf is a broad graph-constrained": (
        "Hydrosheaf M3 supports a controlled public-data capability claim for residence-time emulation, identifiability screening, target withholding, and topology falsification. It does not provide field validation, confirm flow paths, demonstrate universal graph benefit, or validate the complete framework."
    ),
})

MAIN_REPLACEMENTS.update({
    "Table 2. Comparative performance": "Table 2. Active design-matrix performance on the national release. Metrics are reported only for fits passing the reference-free identifiability guard; non-reportable rows are counted but excluded from performance claims. Published LPM outputs are model-derived reference targets, not true ages.",
    "Table 3. Paired ablation analysis": "Table 3. Paired ablation and sensitivity effects on reportable common rows. Negative mean delta denotes lower reference discrepancy than the paired reported-corrected baseline; small supports and shared provenance limit causal interpretation.",
    "Table 4. Graph-constrained benchmark": "Table 4. Graph-age benchmark on the 329 strict reportable nodes. Delta RMSE is graph minus single-node log₁₀ RMSE. A candidate is not a robust improvement if practical within-factor-two agreement deteriorates. Edges are hypothesised connectivity, not verified flow paths.",
    "Table 5. Summary comparison": "Table 5. Modelling-mode comparison on unequal full-available supports and N = 40 common support. Direct ranking is limited to common support. Reported-configuration emulation and fraction sensitivity reproduce model-derived USGS products and are not independent validation.",
    "Full Available =": "Full Available = all reportable metric rows for a mode; Common Support = the same 40 reportable samples shared by compared modes. Full-available rows are descriptive and must not be ranked across unequal supports.",
    "Table 6. Statistical validation": "Table 6. Paired comparison of reported-configuration emulation and reported-output-constrained fraction sensitivity. Statistical differences quantify closer agreement with shared-provenance products; they do not establish independent true-age accuracy.",
    "Figure 2.": "Figure 2. Active design-matrix reference agreement and reportability. Metrics use only identifiable reportable fits; archived withdrawn scenarios are excluded. Lower discrepancy does not imply field-validated true-age accuracy.",
    "Figure 3.": "Figure 3. Strict reported-configuration emulation against model-derived USGS reference ages for the N = 329 reportable subset. This evaluates reference-workflow agreement, not independently known groundwater ages.",
    "Figure 4.": "Figure 4. Graph-age effects on the N = 329 strict reportable nodes. The best candidate RMSE change is numerically small and accompanied by lower within-factor-two agreement; no candidate meets the robust-improvement rule.",
    "Figure 5.": "Figure 5. Age-class reference-discrepancy and young-tracer discordance diagnostics on the strict reportable subset. Age classes derive from model outputs and have unequal support.",
    "Figure 6.": "Figure 6. Leakage-guarded target-tracer withholding. Each target is removed before every single-node and graph-node refit, and only reportable fits are scored in native units. Candidate graphs show no meaningful ³H, SF₆, or ¹⁴C gain.",
    "Figure 7.": "Figure 7. Signed reference-discrepancy diagnostics on the finite strict reportable plotting subset. Counts differ from Supplementary Table S3 because that table classifies old-water evidence before plotting filters.",
})

SUPP_REPLACEMENTS = {
    "A median absolute log₁₀ error": "For the N = 329 strict reportable subset, median absolute log₁₀ reference discrepancy was 0.0279 (multiplicative factor approximately 1.07), log₁₀ RMSE was 0.2769, and 87.5% were within a factor of two. The N = 289 reported-output-constrained fraction sensitivity had median discrepancy 0.0216 and RMSE 0.1964. These quantify agreement with model-derived LPM outputs, not independently known true ages.",
    "Supplementary Table S2.": "Supplementary Table S2. Age-stratified reference-workflow agreement for strict reported-configuration emulation on the reportable subset. Age classes derive from model outputs and supports are unequal.",
    "Supplementary Table S3.": "Supplementary Table S3. Diagnostic ¹⁴C–⁴He agreement categories for all harmonised rows. These categories are not independent age validation and do not restore the withdrawn hierarchical prior.",
    "Supplementary Table S4.": "No Supplementary Table S4 is reported: the final audit found zero supported paired raw and corrected atmospheric-equivalent gas values, so an empirical gas-correction effect could not be estimated.",
    "Supplementary Figure S1.": "No Supplementary Figure S1 is reported because the gas-correction comparison had zero supported paired values after the final audit.",
    "Supplementary Figure S2.": "Supplementary Figure S2. Spatial-distance and vertical-difference distributions for graph edges in the N = 329 strict reportable-node benchmark. These diagnose hypothesised edge geometry, not verified flow connections.",
    "Supplementary Figure S3.": "Supplementary Figure S3. CFC-11 (N = 4) and CFC-12 (N = 6) target-withheld RMSE after the reportability guard. No reportable CFC nodes formed eligible graph edges, so identical bars indicate a non-estimable graph effect, not equivalence.",
}

HIGHLIGHT_REPLACEMENTS = {
    "Age-fraction constraints": "Reported age fractions improved agreement with shared-provenance USGS LPM outputs, not true-age validation.",
    "Graph priors did not": "No candidate graph robustly improved the 329-node reference benchmark or target-withheld prediction.",
    "Physically implausible": "Wrong-direction and randomised controls degraded reference agreement, supporting topology falsification.",
    "Depth-constrained": "Leakage-guarded ³H, SF₆, and ¹⁴C withholding showed no meaningful candidate-graph gain.",
    "Integrated young-": "Only reportable fits are interpreted; CFC graph effects were non-estimable because no edges remained.",
}


def replace_by_prefix(doc: Document, mapping: dict[str, str]) -> None:
    for prefix, replacement in mapping.items():
        matches = [p for p in doc.paragraphs if p.text.strip().startswith(prefix)]
        if len(matches) != 1:
            raise RuntimeError(f"Expected one paragraph starting {prefix!r}; found {len(matches)}")
        matches[0].text = replacement


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def fmt(value) -> str:
    if pd.isna(value):
        return "—"
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, str):
        return value.replace("_", " ")
    value = float(value)
    if math.isclose(value, round(value), abs_tol=1e-12):
        return str(int(round(value)))
    if 0 < abs(value) < 0.001:
        return f"{value:.2e}"
    return f"{value:.4f}"


def replace_table(table, data, headers, columns, font_size=7.0) -> None:
    while len(table.columns) < len(columns):
        table.add_column(Inches(0.6))
    if len(table.columns) != len(columns):
        raise RuntimeError(f"Table has {len(table.columns)} columns; expected {len(columns)}")
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for _ in range(len(data)):
        table.add_row()
    for j, heading in enumerate(headers):
        table.cell(0, j).text = heading
    for i, (_, row) in enumerate(data.iterrows(), start=1):
        for j, column in enumerate(columns):
            table.cell(i, j).text = fmt(row[column])
    repeat_header(table.rows[0])
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)
                    run.font.bold = i == 0


def replace_pictures(doc, paths, remove_first=False) -> None:
    paragraphs = [p for p in doc.paragraphs if p._p.xpath(".//a:blip")]
    expected = len(paths) + int(remove_first)
    if len(paragraphs) != expected:
        raise RuntimeError(f"Expected {expected} picture paragraphs; found {len(paragraphs)}")
    if remove_first:
        paragraphs[0]._p.clear_content()
        paragraphs = paragraphs[1:]
    for paragraph, path in zip(paragraphs, paths, strict=True):
        paragraph._p.clear_content()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(path), width=Inches(6.5))


def update_main() -> None:
    path = DOC_DIR / "Manucript_3.docx"
    doc = Document(path)
    replace_by_prefix(doc, MAIN_REPLACEMENTS)
    specs = [
        (9, "Manuscript_Table1_Nuclear_Tracer_Physics.csv",
         ["Tracer", "Decay/process scale", "Target range (yr)", "Unit", "Benchmark role"],
         ["Tracer", "Decay or process scale", "Target age range (yr)", "Unit", "Benchmark role"], 7.5),
        (10, "Manuscript_Table2_Design_Matrix_Performance.csv",
         ["Scenario", "Total", "Reportable", "Median |log₁₀|", "log₁₀ RMSE", "log₁₀ R²", "Within ×2", "Within ×10", "Calibrated ⁴He"],
         ["scenario_id", "total_rows", "identifiable_rows", "median_abs_log10_error", "log10_rmse", "log10_r2", "within_factor_2", "within_factor_10", "calibrated_he4_rows"], 6.5),
        (11, "Manuscript_Table3_Paired_Ablation_Effects.csv",
         ["Scenario", "Paired N", "Median Δ|log₁₀|", "Mean Δ|log₁₀|", "Improved fraction", "Gained ×2", "Lost ×2"],
         ["scenario_id", "paired_rows", "median_delta_log10_error", "mean_delta_log10_error", "improved_fraction", "gained_factor_2_rows", "lost_factor_2_rows"], 6.7),
        (12, "Manuscript_Table4_Real_USGS_Graph_Benchmark.csv",
         ["Graph", "Type", "Strength", "Nodes", "Edges", "RMSE single", "RMSE graph", "ΔRMSE", "Within ×2 single", "Within ×2 graph", "Violations before", "Violations after", "Improved RMSE"],
         ["graph_family", "control_type", "prior_strength", "n_nodes", "n_edges", "rmse_single_log10", "rmse_graph_log10", "delta_rmse_graph_minus_single", "within_factor_2_single", "within_factor_2_graph", "n_violations_before", "n_violations_after", "improved_vs_single"], 5.8),
        (13, "Manuscript_Table5_Mode_Comparison.csv",
         ["Support", "Mode", "N", "Finite", "Median |log₁₀|", "log₁₀ RMSE", "log₁₀ R²", "Within ×2"],
         ["subset", "mode", "N", "finite_metrics", "median_abs_log10_error", "log10_rmse", "log10_r2", "within_factor_2"], 6.3),
        (14, "Manuscript_Table6_Statistical_Significance.csv",
         ["Test/metric", "Comparison", "Statistic", "p value", "CI lower", "CI upper", "Interpretation"],
         ["test_or_metric", "comparison", "statistic", "p_value", "ci_lower", "ci_upper", "interpretation"], 6.3),
    ]
    for idx, filename, headers, columns, size in specs:
        replace_table(doc.tables[idx], pd.read_csv(TABLE_DIR / filename), headers, columns, size)
    replace_pictures(doc, [
        FIG_DIR / "Manuscript_Fig1_Atmospheric_Histories.png",
        FIG_DIR / "Manuscript_Fig2_Design_Matrix_Performance.png",
        FIG_DIR / "Manuscript_Fig3_USGS_Benchmark_Parity.png",
        FIG_DIR / "Manuscript_Fig4_Real_USGS_Graph_Benchmark.png",
        FIG_DIR / "Manuscript_Fig5_Age_Class_Diagnostics.png",
        FIG_DIR / "Manuscript_Fig6_Cross_Validation_Results.png",
        FIG_DIR / "Manuscript_Fig7_Residual_Diagnostics.png",
    ])
    doc.core_properties.comments = "M3 accuracy lock: reportable fits only; reference agreement is not field truth."
    doc.save(path)


def update_supplement() -> None:
    path = DOC_DIR / "Supplementary_Information_M3.docx"
    doc = Document(path)
    replace_by_prefix(doc, SUPP_REPLACEMENTS)
    for row in doc.tables[7].rows[1:]:
        if "Mississippi River Valley" in row.cells[0].text:
            row.cells[2].text = "Available site/tracer files; reported Table8_LPMs unavailable"
            row.cells[3].text = "Site/tracer context only; no comparable reported-LPM target"
            row.cells[4].text = "Excluded from parity, graph-age, and replication metrics"
            row.cells[5].text = "0 rows in active performance tables"
    s2 = pd.read_csv(TABLE_DIR / "Manuscript_Supp_TableS2_Age_Class_Performance.csv")
    replace_table(doc.tables[8], s2,
                  ["Scenario", "Reference-age class", "N", "Median |log₁₀|", "log₁₀ RMSE", "Within ×2", "Within ×10", "Median proxy coherence"],
                  list(s2.columns), 7.0)
    s3 = pd.read_csv(TABLE_DIR / "Manuscript_Supp_TableS3_Old_Groundwater_Diagnostics.csv")
    replace_table(doc.tables[9], s3,
                  ["Scenario", "Old-water status", "N", "Median ¹⁴C age", "Median ⁴He age", "Median gap log₁₀"],
                  list(s3.columns), 7.0)
    doc.tables[10]._element.getparent().remove(doc.tables[10]._element)
    replace_pictures(doc, [
        FIG_DIR / "Suppl_Fig2_Graph_Edge_Properties.png",
        FIG_DIR / "Suppl_Fig3_CFC_CV_Performance.png",
    ], remove_first=True)
    refs = [p for p in doc.paragraphs if p.text.strip() == "References"]
    if len(refs) != 1:
        raise RuntimeError("Could not locate Supplementary References heading")
    image_p = refs[0].insert_paragraph_before()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.add_run().add_picture(str(FIG_DIR / "Suppl_Fig4_Network_Dating_Ambiguity.png"), width=Inches(6.5))
    caption = refs[0].insert_paragraph_before(
        "Supplementary Figure S4. Controlled-synthetic network-dating ambiguity and recovery diagnostic. "
        "This capability demonstration uses known simulated ages to test implementation behaviour; it is not field validation, "
        "does not establish accuracy for the USGS benchmark, and does not justify a general graph-improvement claim."
    )
    caption.style = refs[0].style
    doc.core_properties.comments = "M3 supplement accuracy lock; withdrawn unsupported gas comparison."
    doc.save(path)


def update_highlights() -> None:
    path = DOC_DIR / "Highlights-AGeochemistry.docx"
    doc = Document(path)
    replace_by_prefix(doc, HIGHLIGHT_REPLACEMENTS)
    doc.save(path)


def update_cover_letter() -> None:
    path = DOC_DIR / "Cover Letter-AGeochemistry.docx"
    doc = Document(path)
    replace_by_prefix(doc, {
        "10 June 2026": "28 July 2026",
        "The principal contribution of the manuscript": (
            "The manuscript presents an explicitly bounded result. Only fits passing a reference-free reportability guard are interpreted. On 329 reportable nodes, no candidate graph robustly improved the reference benchmark: the smallest RMSE change was −0.00135 log₁₀ but within-factor-two agreement deteriorated. Leakage-guarded withholding of ³H, SF₆, and ¹⁴C likewise showed no meaningful candidate-graph gain, while CFC graph effects were non-estimable because no eligible edges remained. Reported age fractions improved agreement with shared-provenance USGS LPM outputs and are therefore a non-independent sensitivity rather than true-age validation. The hierarchical old-water prior and unsupported gas-correction comparison were withdrawn. The contribution is a reproducible framework for reportability screening and topology falsification, not universal predictive superiority."
        ),
    })
    doc.save(path)


def main() -> None:
    update_main()
    update_supplement()
    update_highlights()
    update_cover_letter()
    print("Updated M3 manuscript, supplement, highlights, and cover letter.")


if __name__ == "__main__":
    main()
